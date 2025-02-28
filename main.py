# app.py
import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import numpy as np
from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document

# Function to load data from CSV file
def load_data(uploaded_file):
    return pd.read_csv(uploaded_file)

# Function to plot distribution and return figure
def plot_distribution(df):
    numeric_cols = df.select_dtypes(include=[np.number]).columns
    fig, axes = plt.subplots(len(numeric_cols), 1, figsize=(8, len(numeric_cols) * 4))
    if len(numeric_cols) == 1:
        axes = [axes]  # Ensure axes is iterable
    for i, col in enumerate(numeric_cols):
        sns.histplot(df[col], kde=True, ax=axes[i])
        axes[i].set_title(f'Distribution Plot for {col}')
        axes[i].set_xlabel(col)
        axes[i].set_ylabel('Frequency')
    plt.tight_layout()
    return fig

# Function to plot scatter and return figure
def plot_scatter(df):
    numeric_cols = df.select_dtypes(include=[np.number]).columns
    fig, ax = plt.subplots(figsize=(10, 6))
    sns.pairplot(df[numeric_cols])
    plt.suptitle('Scatter Plots for All Numeric Columns', y=1.02)
    return fig

# Function to plot box plot and return figure
def plot_box(df):
    numeric_cols = df.select_dtypes(include=[np.number]).columns
    fig, axes = plt.subplots(len(numeric_cols), 1, figsize=(8, len(numeric_cols) * 4))
    if len(numeric_cols) == 1:
        axes = [axes]  # Ensure axes is iterable
    for i, col in enumerate(numeric_cols):
        sns.boxplot(x=df[col], ax=axes[i])
        axes[i].set_title(f'Box Plot for {col}')
    plt.tight_layout()
    return fig

# Function to plot relationship plot and return figure
def plot_relationship(df):
    fig = plt.figure(figsize=(8, 6))
    sns.pairplot(df)
    plt.suptitle('Relationship Plots between Numeric Columns', y=1.02)
    return fig

# Function to plot comparison plot and return figure
def plot_comparison(df):
    fig = plt.figure(figsize=(8, 6))
    sns.pairplot(df, kind='reg')
    plt.suptitle('Comparison Plot of Numeric Columns', y=1.02)
    return fig

# Function to save plots to PDF
def save_plots_to_pdf(figures):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    for fig in figures:
        fig.savefig(buffer, format="png")
        c.drawImage(buffer, 0, 0, width=500, height=400)  # Adjust image placement
        c.showPage()
    c.save()
    buffer.seek(0)
    return buffer

# Function to save plots to Word Document
def save_plots_to_word(figures):
    doc = Document()
    for fig in figures:
        img_path = 'temp_plot.png'
        fig.savefig(img_path)
        doc.add_paragraph(f"Plot:")
        doc.add_picture(img_path)
        doc.add_paragraph("\n")
    doc.save("plots.docx")

# Streamlit app
def main():
    st.title('CSV Data Visualizer')

    uploaded_file = st.file_uploader("Upload your CSV file", type=["csv"])
    if uploaded_file is not None:
        # Load and display the dataframe
        df = load_data(uploaded_file)
        st.write(df.head())

        # Generate and display plots
        figures = []

        if st.checkbox('Generate Distribution Plots'):
            fig = plot_distribution(df)
            st.pyplot(fig)
            figures.append(fig)

        if st.checkbox('Generate Scatter Plots'):
            fig = plot_scatter(df)
            st.pyplot(fig)
            figures.append(fig)

        if st.checkbox('Generate Box Plots'):
            fig = plot_box(df)
            st.pyplot(fig)
            figures.append(fig)

        if st.checkbox('Generate Relationship Plots'):
            fig = plot_relationship(df)
            st.pyplot(fig)
            figures.append(fig)

        if st.checkbox('Generate Comparison Plots'):
            fig = plot_comparison(df)
            st.pyplot(fig)
            figures.append(fig)

        # Provide options to download plots as PDF or Word Document
        if st.button('Download as PDF'):
            pdf_buffer = save_plots_to_pdf(figures)
            st.download_button(
                label="Download PDF",
                data=pdf_buffer,
                file_name="plots.pdf",
                mime="application/pdf"
            )

        if st.button('Download as Word Document'):
            save_plots_to_word(figures)
            with open("plots.docx", "rb") as file:
                st.download_button(
                    label="Download Word Document",
                    data=file,
                    file_name="plots.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

if __name__ == "__main__":
    main()
